from pathlib import Path

import pytest
from docx import Document

from app.services.job_description_service import (
    JobDescriptionExtractionError,
    extract_job_description,
)


PDF_FIXTURE = "data/raw/jd/jd001.pdf"


def test_extract_from_txt(tmp_path: Path):
    file = tmp_path / "jd.txt"
    file.write_text(
        "Platform Engineer\n\nBuilds internal tooling.",
        encoding="utf-8",
    )

    text = extract_job_description(file)

    assert "Platform Engineer" in text
    assert "Builds internal tooling." in text


def test_extract_from_pdf():
    text = extract_job_description(PDF_FIXTURE)

    assert len(text) > 0


def test_extract_from_docx(tmp_path: Path):
    file = tmp_path / "jd.docx"

    document = Document()
    document.add_paragraph("Site Reliability Engineer")
    document.add_paragraph("Owns production reliability.")
    document.save(file)

    text = extract_job_description(file)

    assert "Site Reliability Engineer" in text
    assert "Owns production reliability." in text


def test_rejects_unsupported_extension(tmp_path: Path):
    file = tmp_path / "jd.exe"
    file.write_bytes(b"not a document")

    with pytest.raises(JobDescriptionExtractionError):
        extract_job_description(file)


def test_rejects_empty_txt(tmp_path: Path):
    file = tmp_path / "jd.txt"
    file.touch()

    with pytest.raises(JobDescriptionExtractionError):
        extract_job_description(file)


def test_rejects_corrupted_docx(tmp_path: Path):
    file = tmp_path / "jd.docx"
    file.write_bytes(b"definitely not a real docx file")

    with pytest.raises(JobDescriptionExtractionError):
        extract_job_description(file)


def test_rejects_missing_file(tmp_path: Path):
    file = tmp_path / "does_not_exist.txt"

    with pytest.raises(JobDescriptionExtractionError):
        extract_job_description(file)
