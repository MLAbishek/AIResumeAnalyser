from pathlib import Path

import pytest

from app.core.schemas import DocumentType
from app.ingestion.resume_loader import ResumeLoader
from docx import Document
from app.ingestion.resume_loader import ResumeLoader

def test_load_txt_resume(tmp_path: Path):
    file = tmp_path / "resume.txt"

    file.write_text(
        """
        John Doe
        Python Developer

        Skills:
        Python
        FastAPI
        PostgreSQL
        """,
        encoding="utf-8",
    )

    loader = ResumeLoader()

    document = loader.load(file)

    assert document.document_type == DocumentType.RESUME
    assert len(document.pages) == 1

    extracted_text = document.pages[0].text

    assert "John Doe" in extracted_text
    assert "Python Developer" in extracted_text
    assert "FastAPI" in extracted_text


def test_reject_unsupported_resume(tmp_path: Path):
    file = tmp_path / "resume.exe"
    file.write_bytes(b"fake")

    loader = ResumeLoader()

    with pytest.raises(ValueError):
        loader.load(file)


def test_reject_empty_resume(tmp_path: Path):
    file = tmp_path / "resume.txt"
    file.touch()

    loader = ResumeLoader()

    with pytest.raises(ValueError):
        loader.load(file)


def test_reject_missing_resume(tmp_path: Path):
    file = tmp_path / "does_not_exist.txt"

    loader = ResumeLoader()

    with pytest.raises(ValueError):
        loader.load(file)

def test_load_docx_resume(tmp_path: Path):
    file = tmp_path / "resume.docx"

    # Create a new DOCX document
    document = Document()

    document.add_paragraph("John Doe")
    document.add_paragraph("Python Developer")
    document.add_paragraph(
        "Skills: Python, FastAPI, PostgreSQL"
    )

    document.save(file)

    # Load through the application
    loader = ResumeLoader()

    result = loader.load(file)

    assert result.document_type == DocumentType.RESUME
    assert len(result.pages) == 1

    assert "John Doe" in result.pages[0].text
    assert "Python Developer" in result.pages[0].text
    assert "FastAPI" in result.pages[0].text
    assert "PostgreSQL" in result.pages[0].text
    
def test_reject_missing_docx_resume(tmp_path: Path):
    file = tmp_path / "missing.docx"

    loader = ResumeLoader()

    with pytest.raises(ValueError):
        loader.load(file)

def test_reject_corrupted_docx_resume(tmp_path: Path):
    file = tmp_path / "corrupted.docx"

    file.write_bytes(b"This is not a valid DOCX file")

    loader = ResumeLoader()

    with pytest.raises(ValueError):
        loader.load(file)